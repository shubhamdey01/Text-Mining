from pypdf import PdfReader
from docx import Document

def countLinesText(file):
    f = open(file, 'r')
    dat = f.read().split('.')
    f.close()
    dat.pop()
    count = len(dat)
    return count

def countWordsText(file):
    f = open(file, 'r')
    count = 0
    line = f.readline()
    while line:
        count += len(line.split())
        line = f.readline()
    f.close()
    return count

def countCharsText(file):
    f = open(file, 'r')
    count = 0
    line = f.readline()
    while line:
        for word in line.split():
            for i in word:
                if i.isalnum():
                    count += 1
        line = f.readline()
    f.close()
    return count


def countLinesPDF(file):
    f = PdfReader(file)
    count = 0
    for page in f.pages:
        dat = page.extract_text().split('.')
        dat.pop()
        count += len(dat)
    return count

def countWordsPDF(file):
    f = PdfReader(file)
    count = 0
    for page in f.pages:
        dat = page.extract_text().split(' ')
        for i in dat:
            if i != '':
                count += 1
    return count

def countCharsPDF(file):
    f = PdfReader(file)
    count = 0
    for page in f.pages:
        dat = page.extract_text()
        for i in dat:
            if i.isalnum():
                count += 1
    return count


def countLinesDOCX(file):
    f = Document(file)
    count = 0
    for para in f.paragraphs:
        dat = para.text.split('.')
        dat.pop()
        count += len(dat)
    return count

def countWordsDOCX(file):
    f = Document(file)
    count = 0
    for para in f.paragraphs:
        dat = para.text.split(' ')
        for i in dat:
            if i != '':
                count += 1
    return count

def countCharsDOCX(file):
    f = Document(file)
    count = 0
    for para in f.paragraphs:
        dat = para.text
        for i in dat:
            if i.isalnum():
                count += 1
    return count

if __name__ == '__main__':
    # Text File
    print("No. of Sentences = ", countLinesText(r"../Files/data.txt"))
    print("No. of Words = ", countWordsText(r"../Files/data.txt"))
    print("No. of Characters = ", countCharsText(r"../Files/data.txt"))

    # PDF File
    print("\nNo. of Sentences = ", countLinesPDF(r"../Files/data.pdf"))
    print("No. of Words = ", countWordsPDF(r"../Files/data.pdf"))
    print("No. of Characters = ", countCharsPDF(r"../Files/data.pdf"))

    # DOCX File
    print("\nNo. of Sentences = ", countLinesDOCX(r"../Files/data.docx"))
    print("No. of Words = ", countWordsDOCX(r"../Files/data.docx"))
    print("No. of Characters = ", countCharsDOCX(r"../Files/data.docx"))

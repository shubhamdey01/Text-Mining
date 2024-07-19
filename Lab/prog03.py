# read from a '.docx' file and count the number of sentences, words, and characters in the file.

from docx import Document

# count no. of sentences
def countLines(file):
    f = Document(file)
    count = 0
    for para in f.paragraphs:
        dat = para.text.split('.')
        dat.pop()
        count += len(dat)
    return count

# count no. of words
def countWords(file):
    f = Document(file)
    count = 0
    for para in f.paragraphs:
        dat = para.text.split(' ')
        for i in dat:
            if i != '':
                count += 1
    return count

# count no. of characters
def countChars(file):
    f = Document(file)
    count = 0
    for para in f.paragraphs:
        dat = para.text
        for i in dat:
            if i.isalnum():
                count += 1
    return count

print('\nReading from ".txt" file.')
print('No. of Sentences = ', countLines(r'../Files/data.docx'))
print('No. of Words = ', countWords(r'../Files/data.docx'))
print('No. of Characters = ', countChars(r'../Files/data.docx'))